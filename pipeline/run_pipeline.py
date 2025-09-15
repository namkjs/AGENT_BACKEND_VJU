from models.vision_model import VisionModel
from schemas.request_models import ImageAnalysisRequest, ImageAnalysisResponse
from models.agent import Agent
from datetime import datetime
import os
from PIL import Image
import pdf2image
from typing import List
import requests
import tempfile
import urllib.parse
# Thêm import để xử lý DOCX
from docx import Document as DocxDocument
import zipfile
from dotenv import load_dotenv
load_dotenv()
BACKEND_URL = os.getenv("BACKEND_URL")
async def run_full_pipeline(file_url: str) -> dict:
    """Chạy toàn bộ pipeline từ OCR đến Agent approval - hỗ trợ PDF, ảnh, DOCX và MD"""
    
    vision_model = VisionModel()
    agent = Agent()
    
    # Load vision model
    await vision_model.load_model()
    
    # Download file từ URL và lưu tạm
    temp_file_path = await download_file_from_url(file_url)
    
    if not temp_file_path:
        return {
            "approve": "reject",
            "description": "Không thể download file từ URL",
            "ocr_text": "",
            "timestamp": datetime.now().isoformat()
        }
    
    try:
        file_type = get_file_type(temp_file_path)
        print(f"📁 File type detected: {file_type}")
        
        # Xử lý DOCX - không cần OCR
        if file_type == 'docx':
            print("📄 Đang trích xuất text từ file DOCX...")
            docx_text = await extract_text_from_docx(temp_file_path)
            
            if not docx_text.strip():
                return {
                    "approve": "reject", 
                    "description": "File DOCX rỗng hoặc không thể đọc được nội dung",
                    "ocr_text": "",
                    "file_type": file_type,
                    "timestamp": datetime.now().isoformat()
                }
            
            print(f"📚 Đã trích xuất {len(docx_text)} ký tự từ DOCX")
            
            # Gửi text trực tiếp cho Agent
            print("🤖 Đang kiểm tra tài liệu DOCX với Agent...")
            approval_result = await agent.check_document_approval(docx_text.strip())
            
            # Log kết quả
            print(f"📋 Kết quả đánh giá: {approval_result.get('approve', 'unknown')}")
            print(f"💬 Lý do: {approval_result.get('description', 'Không có mô tả')}")
            
            return {
                "approve": approval_result.get("approve", "reject"),
                "description": approval_result.get("description", "Lỗi trong quá trình đánh giá"),
                "ocr_text": docx_text.strip(),
                "total_pages": 1,  # DOCX được coi như 1 "trang"
                "file_type": file_type,
                "file_url": file_url,
                "timestamp": datetime.now().isoformat()
            }
        
        # Xử lý MD - không cần OCR
        elif file_type == 'md':
            print("📄 Đang trích xuất text từ file MD...")
            md_text = await extract_text_from_md(temp_file_path)
            
            if not md_text.strip():
                return {
                    "approve": "reject", 
                    "description": "File MD rỗng hoặc không thể đọc được nội dung",
                    "ocr_text": "",
                    "file_type": file_type,
                    "timestamp": datetime.now().isoformat()
                }
            
            print(f"📚 Đã trích xuất {len(md_text)} ký tự từ MD")
            
            # Gửi text trực tiếp cho Agent
            print("🤖 Đang kiểm tra tài liệu MD với Agent...")
            approval_result = await agent.check_document_approval(md_text.strip())
            
            # Log kết quả
            print(f"📋 Kết quả đánh giá: {approval_result.get('approve', 'unknown')}")
            print(f"💬 Lý do: {approval_result.get('description', 'Không có mô tả')}")
            
            return {
                "approve": approval_result.get("approve", "reject"),
                "description": approval_result.get("description", "Lỗi trong quá trình đánh giá"),
                "ocr_text": md_text.strip(),
                "total_pages": 1,  # MD được coi như 1 "trang"
                "file_type": file_type,
                "file_url": file_url,
                "timestamp": datetime.now().isoformat()
            }
        
        # Xử lý PDF và ảnh - cần OCR như cũ
        else:
            images = await load_file_as_images(temp_file_path)
            
            if not images:
                return {
                    "approve": "reject",
                    "description": "Không thể load file hoặc file không hợp lệ",
                    "ocr_text": "",
                    "timestamp": datetime.now().isoformat()
                }
            
            print(f"📁 Loaded {len(images)} image(s) from: {file_url}")
            
            # Bước 1: OCR từng ảnh
            all_ocr_text = ""
            
            for i, image in enumerate(images):
                print(f"👁️ Đang thực hiện OCR cho ảnh {i+1}/{len(images)}...")
                
                analysis_response: ImageAnalysisResponse = await vision_model.analyze_image(
                    image, 
                    question="Trích xuất toàn bộ text từ tài liệu này một cách chi tiết và chính xác.",
                    max_num=6
                )
                
                # Gộp OCR text của ảnh này - xử lý an toàn
                image_ocr_text = ""
                if hasattr(analysis_response, 'results') and analysis_response.results:
                    for res in analysis_response.results:
                        if hasattr(res, 'description'):
                            image_ocr_text += res.description + "\n"
                        else:
                            image_ocr_text += str(res) + "\n"
                elif isinstance(analysis_response, str):
                    image_ocr_text = analysis_response
                else:
                    image_ocr_text = str(analysis_response)
                
                # Thêm separator giữa các trang
                if image_ocr_text.strip():
                    all_ocr_text += f"--- Trang {i+1} ---\n{image_ocr_text}\n"
                
                print(f"📄 OCR trang {i+1}: {len(image_ocr_text)} ký tự")
            
            print(f"📚 Tổng OCR text: {len(all_ocr_text)} ký tự từ {len(images)} trang")
            
            # Bước 2: Agent kiểm tra approval cho toàn bộ tài liệu
            print("🤖 Đang kiểm tra tài liệu với Agent...")
            approval_result = await agent.check_document_approval(all_ocr_text.strip())
            
            # Log kết quả
            print(f"📋 Kết quả đánh giá: {approval_result.get('approve', 'unknown')}")
            print(f"💬 Lý do: {approval_result.get('description', 'Không có mô tả')}")
            
            return {
                "approve": approval_result.get("approve", "reject"),
                "description": approval_result.get("description", "Lỗi trong quá trình đánh giá"),
                "ocr_text": all_ocr_text.strip(),
                "total_pages": len(images),
                "file_type": get_file_type_from_url(file_url),
                "file_url": file_url,
                "timestamp": datetime.now().isoformat()
            }
            
    except Exception as e:
        print(f"❌ Lỗi trong pipeline: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "approve": "reject",
            "description": f"Lỗi trong quá trình xử lý: {str(e)}",
            "ocr_text": "",
            "timestamp": datetime.now().isoformat()
        }
        
    finally:
        # Cleanup: xóa file tạm
        try:
            if temp_file_path and os.path.exists(temp_file_path):
                os.remove(temp_file_path)
                print(f"🗑️ Đã xóa file tạm: {temp_file_path}")
        except Exception as e:
            print(f"⚠️ Không thể xóa file tạm: {e}")

async def download_file_from_url(file_url: str) -> str:
    """Download file từ URL và lưu vào file tạm"""
    try:
        # Debug original URL
        print(f"🔍 Original file_url: {file_url}")
        
        # URL encode để xử lý ký tự đặc biệt
        encoded_file_path = urllib.parse.quote(file_url, safe='/')
        full_url = f"{BACKEND_URL}/attachments?path={encoded_file_path}"
        print(f"📥 Đang download file từ: {full_url}")
        
        # Gửi request để download file
        response = requests.get(full_url, stream=True, timeout=30)
        
        # Debug response
        print(f"📊 Response status: {response.status_code}")
        if response.status_code == 404:
            print(f"🔍 File not found. Trying alternative path...")
            # Thử path khác nếu cần
            alternative_url = f"{BACKEND_URL}/{file_url}"
            print(f"🔄 Trying: {alternative_url}")
            response = requests.get(alternative_url, stream=True, timeout=30)
        
        response.raise_for_status()
        
        # Lấy extension từ URL hoặc Content-Type
        file_extension = get_file_extension_from_url(file_url, response)
        
        # Tạo file tạm
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            temp_file_path = temp_file.name
            
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    temp_file.write(chunk)
        
        print(f"✅ Download thành công: {temp_file_path}")
        return temp_file_path
        
    except requests.exceptions.RequestException as e:
        print(f"❌ Lỗi download file: {e}")
        return None
    except Exception as e:
        print(f"❌ Lỗi không xác định khi download: {e}")
        return None

def get_file_extension_from_url(file_url: str, response: requests.Response) -> str:
    """Lấy file extension từ URL hoặc Content-Type"""
    try:
        # Debug info
        print(f"🔍 Getting extension for URL: {file_url}")
        
        # Thử lấy từ URL trước
        url_extension = os.path.splitext(file_url.split('?')[0])[1].lower()
        print(f"🔍 URL extension: {url_extension}")
        
        if url_extension in ['.pdf', '.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp', '.docx', '.doc', '.md']:
            return url_extension
        
        # Nếu không có từ URL, thử từ Content-Type
        content_type = response.headers.get('content-type', '').lower()
        print(f"🔍 Content-Type: {content_type}")
        
        if 'pdf' in content_type:
            return '.pdf'
        elif 'jpeg' in content_type:
            return '.jpg'
        elif 'png' in content_type:
            return '.png'
        elif 'bmp' in content_type:
            return '.bmp'
        elif 'tiff' in content_type:
            return '.tiff'
        elif 'webp' in content_type:
            return '.webp'
        elif 'wordprocessingml' in content_type or 'msword' in content_type:
            return '.docx'
        elif 'markdown' in content_type or 'text/plain' in content_type:
            return '.md'
        
        # Default
        return '.pdf'
        
    except Exception as e:
        print(f"⚠️ Không thể xác định file extension: {e}")
        return '.pdf'

async def load_file_as_images(file_path: str) -> List[Image.Image]:
    """Load file thành list các ảnh PIL - hỗ trợ PDF và các định dạng ảnh"""
    try:
        if not os.path.exists(file_path):
            print(f"❌ File không tồn tại: {file_path}")
            return []
        
        normalized_path = os.path.normpath(file_path).replace('\\', '/')
        print(f"🔧 Normalized path: {normalized_path}")
        
        file_extension = os.path.splitext(normalized_path)[1].lower()
        
        if file_extension == '.pdf':
            print("📄 Đang convert PDF thành ảnh...")
            try:
                print(f"🔍 Đường dẫn PDF: {normalized_path}")
                images = pdf2image.convert_from_path(
                    pdf_path=normalized_path, 
                    dpi=300
                )
                print(f"✅ Convert PDF thành công: {len(images)} trang")
                return images
            except Exception as e:
                print(f"❌ Lỗi convert PDF: {str(e)}")
                print(f"🔍 Chi tiết: Đảm bảo Poppler đã cài đặt và PDF hợp lệ. Path: {normalized_path}")
                return []
                
        elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp']:
            # Xử lý ảnh
            print(f"🖼️ Đang load ảnh: {file_extension}")
            try:
                image = Image.open(normalized_path).convert("RGB")
                return [image]
            except Exception as e:
                print(f"❌ Lỗi load ảnh: {str(e)}")
                return []
        else:
            print(f"❌ Định dạng file không được hỗ trợ: {file_extension}")
            return []
            
    except Exception as e:
        print(f"❌ Lỗi load file: {str(e)}")
        return []

async def extract_text_from_docx(file_path: str) -> str:
    """Trích xuất text từ file DOCX"""
    try:
        # Kiểm tra file có phải DOCX hợp lệ không
        if not zipfile.is_zipfile(file_path):
            print("❌ File không phải là DOCX hợp lệ")
            return ""
        
        # Load DOCX document
        doc = DocxDocument(file_path)
        
        # Trích xuất text từ các paragraph
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # Trích xuất text từ tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        
        # Kết hợp tất cả text
        extracted_text = "\n".join(full_text)
        
        print(f"📄 Đã trích xuất {len(extracted_text)} ký tự từ DOCX")
        print(f"📄 Có {len(doc.paragraphs)} paragraphs và {len(doc.tables)} tables")
        
        return extracted_text
        
    except Exception as e:
        print(f"❌ Lỗi khi đọc file DOCX: {str(e)}")
        return ""

async def extract_text_from_md(file_path: str) -> str:
    """Trích xuất text từ file MD"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        print(f"📄 Đã trích xuất {len(content)} ký tự từ MD")
        return content
        
    except Exception as e:
        print(f"❌ Lỗi khi đọc file MD: {str(e)}")
        return ""

def get_file_type_from_url(file_url: str) -> str:
    """Lấy loại file từ URL"""
    try:
        file_extension = os.path.splitext(file_url.split('?')[0])[1].lower()
        if file_extension == '.pdf':
            return 'pdf'
        elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp']:
            return 'image'
        elif file_extension in ['.docx', '.doc']:
            return 'docx'
        elif file_extension == '.md':
            return 'md'
        else:
            return 'unknown'
    except:
        return 'unknown'

def get_file_type(file_path: str) -> str:
    """Lấy loại file từ đường dẫn local"""
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.pdf':
        return 'pdf'
    elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp']:
        return 'image'
    elif file_extension in ['.docx', '.doc']:
        return 'docx'
    elif file_extension == '.md':
        return 'md'
    else:
        return 'unknown'

